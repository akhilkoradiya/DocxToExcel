import docx
import pandas as pd
import os

def extract_headings_from_docx(docx_file):
    doc = docx.Document(docx_file)
    headings = []

    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1':
            headings.append(paragraph.text)

    return headings

def extract_content_between_headers(docx_file, start_header, end_header):
    doc = docx.Document(docx_file)
    content = []
    start_found = False
    current_content = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in start_header:
            start_found = True
            current_content = []
        elif start_found:
            if paragraph.text.strip() in end_header:
                content.append('\n'.join(current_content))
                start_found = False
            else:
                current_content.append(paragraph.text)

    return content

if __name__ == "__main__":
    input_directory = "Data"
    combined_data = []
    for filename in os.listdir("Data"):
        if filename.endswith(".docx"):
            input_docx_file = os.path.join(input_directory, filename)
            extracted_headings = extract_headings_from_docx(input_docx_file)
            extracted_descriptions = extract_content_between_headers(input_docx_file, "Description", "Impact")
            extracted_impacts = extract_content_between_headers(input_docx_file, "Impact", ["Affected Application", "Affected URL"])
            extracted_remediations = extract_content_between_headers(input_docx_file, "Remediation", "References:")
            extracted_references = extract_content_between_headers(input_docx_file, "References:", "Proof of Concept")
            if extracted_headings and extracted_descriptions and extracted_impacts and extracted_remediations and extracted_references:
                for title, description, impact, remediation, references in zip(
                    extracted_headings, 
                    extracted_descriptions, 
                    extracted_impacts, 
                    extracted_remediations,
                    extracted_references
                    ):
                    combined_data.append({
                        "report": os.path.splitext(filename)[0],
                        "title": title,
                        "description": description,
                        "impact": impact,
                        "remediation": remediation,
                        "references" : references
                    })
    df = pd.DataFrame(combined_data)
    with pd.ExcelWriter('output.xlsx', engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Sheet1', index=False)
    print("Excel file has been created.")